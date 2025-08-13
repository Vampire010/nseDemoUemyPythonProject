import requests
from openpyxl import load_workbook, Workbook
from docx import Document
import json
# Input Excel file
excel_file = "udemy_courses.xlsx"
wb = load_workbook(excel_file)
ws = wb.active

# Output Word document setup
document = Document()
document.add_heading("Udemy Course Curriculums", level=1)

# Output Excel workbook setup for lecture data
out_wb = Workbook()
out_ws = out_wb.active
out_ws.title = "Lecture Details"
out_ws.append(["Course ID", "Course Name", "Lecture ID", "Lecture Title", "Object Index"])

# Cookies and headers setup

# === Load access token from Authentication.json ===
with open("Authentication.json", "r") as f:
    auth_data = json.load(f)
ACCESS_TOKEN = auth_data.get("access_token")
CLIENT_ID = auth_data.get("client_id")
CSRF = auth_data.get("csrf")

cookies = {
    "access_token": ACCESS_TOKEN,
    "client_id": CLIENT_ID,
    "csrf": CSRF
}

headers = {
    'Accept': 'application/json, text/plain, */*',
    'X-Requested-With': 'XMLHttpRequest',
    'Cookie': '; '.join([f"{key}={value}" for key, value in cookies.items()])
}

# Iterate over course IDs in Excel
for row in range(2, ws.max_row + 1):
    course_id = ws[f"B{row}"].value
    course_name = ws[f"C{row}"].value
    if not course_id:
        continue

    print(f"📘 Fetching Course ID: {course_id} - {course_name}")
    document.add_page_break()
    document.add_heading(f"Course ID: {course_id} - {course_name}", level=2)

    # Build API URL
    url = f"https://www.udemy.com/api-2.0/courses/{course_id}/subscriber-curriculum-items/?" \
          "curriculum_types=chapter,lecture,practice,quiz,role-play&page_size=200" \
          "&fields[lecture]=id,title,object_index,is_published,sort_order,created,asset,supplementary_assets,is_free" \
          "&fields[quiz]=title,object_index,is_published,sort_order,type" \
          "&fields[practice]=title,object_index,is_published,sort_order" \
          "&fields[chapter]=title,object_index,is_published,sort_order" \
          "&fields[asset]=title,filename,asset_type,status,time_estimation,is_external" \
          "&caching_intent=True"

    response = requests.get(url, headers=headers)
    if response.status_code != 200:
        document.add_paragraph(f"❌ Failed to fetch course. Status: {response.status_code}")
        continue

    data = response.json()
    section_count = 0
    lecture_count = 0

    for item in data.get('results', []):
        if item['_class'] == 'chapter':
            section_count += 1
            object_index = item.get("object_index", "N/A")
            section_title = f"Section {section_count:02d} - {item['title']} (Index: {object_index})"
            document.add_heading(section_title, level=3)

        elif item['_class'] == 'lecture':
            lecture_count += 1
            lecture_id = item.get("id")
            lecture_title = item.get("title", "Untitled")
            object_index = item.get("object_index", "N/A")

            time_seconds = item.get('asset', {}).get('time_estimation')
            time_minutes = round(time_seconds / 60) if time_seconds else "N/A"
            time_str = f"{time_minutes} min" if time_seconds else "N/A"

            # Write to Word
            document.add_paragraph(f"{lecture_count:02d} - {lecture_title} (Time: {time_str})")

            # Write to Excel
            out_ws.append([course_id, course_name, lecture_id, lecture_title, object_index])

            # Resources
            if item.get('supplementary_assets'):
                for asset in item['supplementary_assets']:
                    resource_title = asset.get('title', 'Untitled Resource')
                    document.add_paragraph(f"Resource: {resource_title}")

# Save Word and Excel files
print("\n💾 Saving Word and Excel files...")
document.save("all_udemy_courses_curriculum.docx")
out_wb.save("udemy_lecture_details.xlsx")

print("✅ Word file saved: all_udemy_courses_curriculum.docx")
print("✅ Excel file saved: udemy_lecture_details.xlsx")
