import requests
from openpyxl import load_workbook, Workbook
from docx import Document

# Load input Excel file
input_excel_file = "udemy_lecture_details.xlsx"
wb = load_workbook(input_excel_file)
ws = wb.active

# Word document setup
document = Document()
document.add_heading("Udemy Course Curriculums", level=1)

# Excel output setup
output_wb = Workbook()
output_ws = output_wb.active
output_ws.title = "Lecture Info"
output_ws.append(["Course ID", "Course Name", "Lecture ID", "Lecture Title", "Object Index", "Time (min)"])

# Headers and Cookies
cookies = {
    "access_token": "qRD3K7gMIf+phY7p+ZWfqlNkVKSYv4en9+SGNu+S04Y:g+C7av1pzqNxwxmJPEPQMGH9ebqHRxNR6QG8nSy8xHw",
    "client_id": "bd2565cb7b0c313f5e9bae44961e8db2",
    "csrf": "cJ0qnLjwfnaYxnxuntf7AbwC86JQGxy0"
}

headers = {
    'Accept': 'application/json, text/plain, */*',
    'X-Requested-With': 'XMLHttpRequest',
    'Cookie': '; '.join([f"{key}={value}" for key, value in cookies.items()])
}

# Iterate over input Excel rows (starting from row 2)
for row in range(2, ws.max_row + 1):
    course_id = ws[f"A{row}"].value
    course_name = ws[f"C{row}"].value or ""
    if not course_id:
        continue

    print(f"📘 Fetching Course ID: {course_id} - {course_name}")
    document.add_page_break()
    document.add_heading(f"{course_name} ({course_id})", level=2)

    url = f"https://www.udemy.com/api-2.0/courses/{course_id}/subscriber-curriculum-items/?" \
          "curriculum_types=chapter,lecture,practice,quiz,role-play&page_size=200" \
          "&fields[lecture]=id,title,object_index,is_published,sort_order,created,asset,supplementary_assets,is_free" \
          "&fields[quiz]=title,object_index,is_published,sort_order,type" \
          "&fields[practice]=title,object_index,is_published,sort_order" \
          "&fields[chapter]=title,object_index,is_published,sort_order" \
          "&fields[asset]=title,filename,asset_type,status,time_estimation,is_external"

    response = requests.get(url, headers=headers)
    if response.status_code != 200:
        print(f"❌ Failed to fetch course. Status: {response.status_code}")
        document.add_paragraph(f"❌ Failed to fetch course. Status: {response.status_code}")
        continue

    data = response.json()
    section_count = 0
    lecture_count = 0  # Continuous across entire course

    for item in data.get('results', []):
        if item['_class'] == 'chapter':
            section_count += 1
            section_title = f"Section {section_count:02d} - {item.get('title', 'Untitled')}"
            document.add_heading(section_title, level=3)

        elif item['_class'] == 'lecture':
            lecture_count += 1
            lecture_id = item.get("id")
            lecture_title = item.get("title", "Untitled")
            object_index = item.get("object_index", "N/A")

            # Time in minutes
            seconds = item.get('asset', {}).get('time_estimation')
            time_minutes = round(seconds / 60) if isinstance(seconds, (int, float)) else "N/A"

            # Word output
            lecture_line = f"{lecture_count:02d} - {lecture_title} (Index: {object_index}, Time: {time_minutes} min)"
            document.add_paragraph(lecture_line, style='List Bullet')

            # Excel output
            output_ws.append([course_id, course_name, lecture_id, lecture_title, object_index, time_minutes])

            # Resources
            if item.get('supplementary_assets'):
                for asset in item['supplementary_assets']:
                    resource_title = asset.get('title', 'Untitled Resource')
                    document.add_paragraph(f"Resource: {resource_title}", style='List Bullet 2')

# Save files
document.save("all_udemy_courses_curriculum.docx")
output_wb.save("all_lecture_data.xlsx")

print("✅ Word saved: all_udemy_courses_curriculum.docx")
print("✅ Excel saved: all_lecture_data.xlsx")
